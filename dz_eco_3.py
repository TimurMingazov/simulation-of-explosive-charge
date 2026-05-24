import io
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document

# ---------- Данные для задания 3.1 (X в зависимости от варианта) ----------
x_values_31 = {
    1: 110, 2: 115, 3: 120, 4: 125, 5: 130,
    6: 135, 7: 140, 8: 145, 9: 150, 10: 155
}

# ---------- Данные для задания 3.2 (X в зависимости от варианта) ----------
x_values_32 = {
    1: 980, 2: 982, 3: 984, 4: 986, 5: 988,
    6: 990, 7: 992, 8: 994, 9: 996, 10: 998
}

TEMPLATE_PATH = Path(__file__).with_name("pattern_economika_dz3.docx")


def fmt(value, digits=2):
    """Формат числа для отчёта: 1234.50 -> 1234,50."""
    return f"{value:.{digits}f}".replace(".", ",")


def calculate_report_data(variant: int, student_name: str = "") -> dict:
    """Считает все показатели и возвращает словарь для подстановки в DOCX-шаблон."""
    x31 = x_values_31[variant]
    x32 = x_values_32[variant]

    # Задание 3.1
    total_cost_31 = 200 + 140 + 80 + 80 + 60 + 100 + 40 + x31 + 200
    indirect = 80 + 80 + 60 + 40 + x31 + 200
    base_A, base_B = 120, 80
    k_A, k_B = base_A / 200, base_B / 200
    indirect_A, indirect_B = indirect * k_A, indirect * k_B
    direct_A, direct_B = 120 + 80 + 40, 80 + 60 + 60
    full_A, full_B = direct_A + indirect_A, direct_B + indirect_B
    unit_A, unit_B = full_A / 200 * 1000, full_B / 400 * 1000

    # Задание 3.2
    revenue = x32 * 1000 / 1000  # тыс. руб.
    total_cost_32 = 250 + 150 + 160 + 140
    profit_sales = revenue - total_cost_32
    profit_gross = profit_sales + 50 - 10
    tax = profit_gross * 0.24
    profit_net = profit_gross - tax
    prod_funds = 600 + 200
    rent_company = (profit_gross / prod_funds) * 100
    rent_product = (profit_sales / total_cost_32) * 100
    economic_profit = revenue - (total_cost_32 + 10 + 50 + 100 + 200 * 0.18)

    if economic_profit > 0:
        economic_profit_text = f"Предприятие создавать целесообразно, экономическая прибыль положительна ({fmt(economic_profit)} тыс. руб.)"
    else:
        economic_profit_text = f"Предприятие создавать нецелесообразно, экономическая прибыль отрицательна ({fmt(economic_profit)} тыс. руб.)"

    return {
        "var": str(variant),
        "name": student_name.strip() or "________________",

        # В шаблоне встречаются и латинская X, и кириллическая Х, поэтому заполняем оба варианта.
        "X31": str(x31),
        "Х31": str(x31),
        "X32": str(x32),
        "Х32": str(x32),

        "total_cost": fmt(total_cost_31),
        "indirect": fmt(indirect),
        "k_A": fmt(k_A, 1),
        "k_B": fmt(k_B, 1),
        "indirect_A": fmt(indirect_A),
        "indirect_B": fmt(indirect_B),
        "full_A": fmt(full_A),
        "full_B": fmt(full_B),
        "unit_A": fmt(unit_A),
        "unit_B": fmt(unit_B),

        "revenue": fmt(revenue),
        "profit_sales": fmt(profit_sales),
        "profit_gross": fmt(profit_gross),
        "profit_net": fmt(profit_net),
        "rent_company": fmt(rent_company),
        "rent_product": fmt(rent_product),
        "economic_profit": fmt(economic_profit),
        "economic_profit_v": economic_profit_text,
    }


def replace_in_paragraph(paragraph, replacements: dict):
    """Заменяет {{метки}} в абзаце. Сначала пытается сохранить runs, затем fallback."""
    changed = False

    # Быстрый путь: метка целиком находится в одном run.
    for run in paragraph.runs:
        for key, value in replacements.items():
            marker = "{{" + key + "}}"
            if marker in run.text:
                run.text = run.text.replace(marker, str(value))
                changed = True

    # Fallback: Word иногда разбивает {{marker}} на несколько runs.
    full_text = paragraph.text
    new_text = full_text
    for key, value in replacements.items():
        marker = "{{" + key + "}}"
        new_text = new_text.replace(marker, str(value))

    if new_text != full_text:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)
        changed = True

    return changed


def iter_all_paragraphs(document: Document):
    """Абзацы в документе, таблицах, header/footer."""
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def build_docx_report(variant: int, student_name: str) -> io.BytesIO:
    """Создаёт DOCX-отчёт из шаблона pattern_economika_dz3.docx."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Не найден шаблон {TEMPLATE_PATH.name}. Положите его рядом с этим .py файлом."
        )

    document = Document(TEMPLATE_PATH)
    replacements = calculate_report_data(variant, student_name)

    for paragraph in iter_all_paragraphs(document):
        replace_in_paragraph(paragraph, replacements)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def show_report_31(variant):
    X = x_values_31[variant]
    data = calculate_report_data(variant)

    indirect = float(data["indirect"].replace(",", "."))
    indirect_A = float(data["indirect_A"].replace(",", "."))
    indirect_B = float(data["indirect_B"].replace(",", "."))
    full_A = float(data["full_A"].replace(",", "."))
    full_B = float(data["full_B"].replace(",", "."))
    unit_A = float(data["unit_A"].replace(",", "."))
    unit_B = float(data["unit_B"].replace(",", "."))

    st.header("📌 ЗАДАНИЕ 3.1")
    st.markdown("---")

    with st.expander("📊 1. Смета затрат на производство", expanded=True):
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("**Формула расчёта:**")
            st.latex(r"S = \sum_{i=1}^{9} 3_i")
            st.latex(r"S = 3\Pi_{пр} + M + 3\Pi_{АУ} + 3\Pi_{всп} + A_{зд} + Э_{тех} + Э_{осв} + A_{об} + 3_{пр}")
        with col2:
            st.markdown("**Подстановка:**")
            st.latex(f"S = 200 + 140 + 80 + 80 + 60 + 100 + 40 + {X} + 200")
            st.latex(f"S = {data['total_cost']} \\text{{ тыс. руб.}}")

    with st.expander("📈 2. Распределение косвенных расходов", expanded=True):
        st.markdown("**База распределения** – заработная плата производственных рабочих:")
        st.latex(r"B = 3\Pi_A + 3\Pi_B = 120 + 80 = 200 \text{ тыс. руб.}")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Коэффициенты:**")
            st.latex(r"k_A = \frac{120}{200} = 0,6")
            st.latex(r"k_B = \frac{80}{200} = 0,4")
        with col2:
            st.markdown("**Косвенные расходы:**")
            st.latex(r"\text{КР} = 80+80+60+40+" + str(X) + r"+200 = " + f"{indirect:.2f}")
            st.latex(r"\text{КР}_A = " + f"{indirect:.2f} \\times 0,6 = {indirect_A:.2f}")
            st.latex(r"\text{КР}_B = " + f"{indirect:.2f} \\times 0,4 = {indirect_B:.2f}")

    with st.expander("📋 3. Калькуляция себестоимости продукции", expanded=True):
        calc_data = {
            "Статья затрат": ["Прямые затраты", "Косвенные расходы", "Полная себестоимость", "Количество", "Себестоимость единицы"],
            "Изделие А": ["240,00 тыс. руб.", f"{fmt(indirect_A)} тыс. руб.", f"{fmt(full_A)} тыс. руб.", "200 шт.", f"{fmt(unit_A)} руб./шт."],
            "Изделие Б": ["200,00 тыс. руб.", f"{fmt(indirect_B)} тыс. руб.", f"{fmt(full_B)} тыс. руб.", "400 шт.", f"{fmt(unit_B)} руб./шт."],
        }
        df = pd.DataFrame(calc_data).set_index("Статья затрат")
        st.dataframe(df, use_container_width=True)
        st.success(f"✅ **Результат:** Себестоимость единицы А = {fmt(unit_A)} руб., Б = {fmt(unit_B)} руб.")


def show_report_32(variant):
    X = x_values_32[variant]
    data = calculate_report_data(variant)

    revenue = float(data["revenue"].replace(",", "."))
    total_cost = 700
    profit_sales = float(data["profit_sales"].replace(",", "."))
    profit_gross = float(data["profit_gross"].replace(",", "."))
    profit_net = float(data["profit_net"].replace(",", "."))
    rent_company = float(data["rent_company"].replace(",", "."))
    rent_product = float(data["rent_product"].replace(",", "."))
    economic_profit = float(data["economic_profit"].replace(",", "."))
    tax = profit_gross * 0.24

    st.header("📌 ЗАДАНИЕ 3.2")
    st.markdown("---")

    with st.expander("💰 1. Расчёт прибыли", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Выручка:**")
            st.latex(r"B = \text{Ц} \times Q = " + str(X) + r" \times 1000 = " + f"{revenue:.2f} \\text{{ тыс. руб.}}")
            st.markdown("**Прибыль от реализации:**")
            st.latex(r"\Pi_{реал} = B - С = " + f"{revenue:.2f} - {total_cost:.2f} = {profit_sales:.2f}")
            st.markdown("**Валовая прибыль:**")
            st.latex(r"\Pi_{вал} = \Pi_{реал} + Д_{проч} - \%_{кред} = " + f"{profit_sales:.2f} + 50 - 10 = {profit_gross:.2f}")
        with col2:
            st.markdown("**Налог на прибыль (24%):**")
            st.latex(r"\text{Н} = \Pi_{вал} \times 0,24 = " + f"{profit_gross:.2f} \\times 0,24 = {tax:.2f}")
            st.markdown("**Чистая прибыль:**")
            st.latex(r"\Pi_{чист} = \Pi_{вал} - \text{Н} = " + f"{profit_gross:.2f} - {tax:.2f} = {profit_net:.2f}")

    with st.expander("📊 2. Расчёт рентабельности", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Рентабельность предприятия:**")
            st.latex(r"R_{пред} = \frac{\Pi_{вал}}{ОС + ОБ} \times 100\%")
            st.latex(r"R_{пред} = \frac{" + f"{profit_gross:.2f}" + r"}{600 + 200} \times 100\% = " + f"{rent_company:.2f}\\%")
        with col2:
            st.markdown("**Рентабельность продукции:**")
            st.latex(r"R_{прод} = \frac{\Pi_{реал}}{С} \times 100\%")
            st.latex(r"R_{прод} = \frac{" + f"{profit_sales:.2f}" + r"}{" + f"{total_cost:.2f}" + r"} \times 100\% = " + f"{rent_product:.2f}\\%")

    with st.expander("⚖️ 3. Оценка целесообразности создания предприятия", expanded=True):
        st.markdown("**Экономическая прибыль = выручка − (явные издержки + неявные издержки)**")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Явные издержки:**")
            st.latex(r"\text{ЯИ} = С + \%_{кред} = 700 + 10 = 710")
            st.markdown("**Неявные издержки:**")
            st.latex(r"\text{НИ} = 50 + 100 + 200 \times 0,18 = 186")
        with col2:
            st.markdown("**Экономическая прибыль:**")
            st.latex(r"\Pi_{экон} = B - (\text{ЯИ} + \text{НИ})")
            st.latex(r"\Pi_{экон} = " + f"{revenue:.2f} - (710 + 186) = {economic_profit:.2f}")

        if economic_profit > 0:
            st.success(f"✅ **ВЫВОД:** Предприятие создавать целесообразно, экономическая прибыль положительна ({fmt(economic_profit)} тыс. руб.)")
        else:
            st.error(f"❌ **ВЫВОД:** Предприятие создавать нецелесообразно, экономическая прибыль отрицательна ({fmt(economic_profit)} тыс. руб.)")


def run():
    st.set_page_config(page_title="Экономика предприятия", layout="wide")
    st.title("🏭 Отчёт по экономике предприятия")
    st.markdown("## Тема: Издержки, прибыль и рентабельность")

    st.video("mem.mp4")

    with st.sidebar:
        variant = st.selectbox("Выберите вариант:", list(range(1, 11)))
        student_name = st.text_input("Фамилия И.О. / группа:", placeholder="Например: Иванов И.И., Б123")
        st.markdown("---")
        st.info(f"📌 Выбран вариант {variant}")

        report_bytes = build_docx_report(variant, student_name)
        st.download_button(
            label="📄 Скачать отчёт DOCX",
            data=report_bytes,
            file_name=f"economika_dz3_variant_{variant}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    if st.sidebar.button("🚀 Показать расчёты на сайте", use_container_width=True):
        st.markdown("---")
        show_report_31(variant)
        st.markdown("---")
        show_report_32(variant)
        st.balloons()
    else:
        st.info("Выберите вариант слева и нажмите «Скачать отчёт DOCX» или «Показать расчёты на сайте».")
